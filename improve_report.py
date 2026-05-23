"""Build a final defense-friendly report from generated lab artifacts."""
from __future__ import annotations

import ast
import csv
import hashlib
import json
import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).parent
OUTPUT = ROOT / "report_v16_final.docx"
FORMULA_DIR = ROOT / "formula_images" / "final"
FONT = "Times New Roman"
MONO = "Consolas"
TEXT_SIZE = Pt(14)

plt.rcParams.update({"mathtext.fontset": "stix", "font.family": "STIXGeneral"})


def clean_text(text: str) -> str:
    return text.replace("—", "-").replace("–", "-").replace("−", "-")


def set_run(run, *, mono=False, size=TEXT_SIZE, bold=False, italic=False) -> None:
    name = MONO if mono else FONT
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def add_p(doc, text="", *, align=None, indent=True, bold=False, italic=False, size=TEXT_SIZE):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run(clean_text(text))
        set_run(run, bold=bold, italic=italic, size=size)
    return p


def add_heading(doc, text, *, level=1):
    p = add_p(doc, clean_text(text), bold=True)
    p.paragraph_format.space_before = Pt(8 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_dash(doc, text, *, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run("- " + clean_text(text))
    set_run(run, bold=bold)
    return p


def add_code_block(doc, code: str, *, size=Pt(8.2)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    p_pr.append(shd)
    for i, line in enumerate(clean_text(code).splitlines()):
        run = p.add_run(line)
        set_run(run, mono=True, size=size)
        if i < len(code.splitlines()) - 1:
            run.add_break()
    return p


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = TEXT_SIZE
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def render_formula(latex: str, *, fontsize=15) -> Path:
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha1(f"final-{fontsize}-{latex}".encode("utf-8")).hexdigest()[:16]
    path = FORMULA_DIR / f"formula_{digest}.png"
    if path.exists():
        return path
    fig = plt.figure(figsize=(8.0, 1.0), dpi=300)
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize)
    fig.savefig(path, transparent=True, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return path


def add_formula(doc, latex: str, number: str | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    formula_path = render_formula(latex)
    if number:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        run = p.add_run()
        add_formula_picture(run, formula_path)
        p.add_run("\t")
        num = p.add_run(f"({number})")
        set_run(num)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        add_formula_picture(run, formula_path)
    return p


def add_formula_picture(run, formula_path: Path) -> None:
    with Image.open(formula_path) as img:
        dpi = img.info.get("dpi", (300, 300))[0] or 300
        native_width_cm = img.size[0] / dpi * 2.54
    max_width = Cm(11.5)
    if native_width_cm > max_width.cm:
        run.add_picture(str(formula_path), width=max_width)
    else:
        run.add_picture(str(formula_path))


def add_table_caption(doc, caption: str):
    p = add_p(doc, caption, align="left", indent=False)
    p.paragraph_format.space_before = Pt(4)
    return p


def add_table(doc, headers, rows, *, font_size=Pt(10.5)):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(clean_text(str(header)))
        set_run(run, bold=True, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(clean_text(str(value)))
            set_run(run, size=font_size)
    return table


def read_csv(path: Path):
    with path.open(encoding="utf-8", newline="") as file:
        rows = list(csv.reader(file))
    return rows[0], rows[1:]


def add_figure(doc, path: str, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(ROOT / path), width=Cm(14.2))
    cap = add_p(doc, clean_text(caption), align="center", indent=False, italic=True)
    cap.paragraph_format.space_after = Pt(4)


def extract_function_source(file_path: Path, func_name: str) -> str:
    source = file_path.read_text(encoding="utf-8")
    tree = ast.parse(source)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == func_name:
            lines = source.splitlines()[node.lineno - 1 : node.end_lineno]
            return "\n".join(lines)
    raise ValueError(f"{func_name} not found in {file_path}")


def signature_block(name: str) -> str:
    blocks = {
        "nelder_mead": '''def nelder_mead(func, x0, step=1.0, tol=1e-4, max_iter=300,
                alpha=1.0, gamma=2.0, beta=0.5, sigma=0.5) -> NelderMeadResult:
    """Минимизация функции двух переменных методом Нелдера-Мида."""
    # полная реализация приведена в Приложении А и в репозитории''',
        "bisection_min": '''def bisection_min(func, a: float, b: float, tol: float = 1e-4) -> tuple[float, float]:
    """Одномерная минимизация методом деления отрезка пополам."""
    # полная реализация приведена в Приложении А и в репозитории''',
        "hooke_jeeves": '''def hooke_jeeves(func, x0, delta=1.0, tol=1e-4, max_iter=300) -> HookeJeevesResult:
    """Минимизация функции двух переменных методом Хука-Дживса."""
    # полная реализация приведена в Приложении А и в репозитории''',
        "steepest_descent": '''def steepest_descent(func, grad, q_matrix, x0, tol=1e-4, max_iter=300) -> GradientResult:
    """Минимизация квадратичной функции методом наискорейшего спуска."""
    # полная реализация приведена в Приложении А и в репозитории''',
        "fletcher_reeves": '''def fletcher_reeves(func, grad, q_matrix, x0, tol=1e-4, max_iter=300) -> GradientResult:
    """Минимизация квадратичной функции методом Флетчера-Ривса."""
    # полная реализация приведена в Приложении А и в репозитории''',
    }
    return blocks[name]


def load_values():
    return json.loads((ROOT / "values.json").read_text(encoding="utf-8"))


def add_title(doc):
    add_p(doc, "Министерство образования и науки Российской Федерации", align="center", indent=False)
    add_p(doc, "Федеральное государственное бюджетное образовательное учреждение высшего образования", align="center", indent=False)
    add_p(doc, "«Ижевский государственный технический университет имени М.Т. Калашникова»", align="center", indent=False)
    for _ in range(5):
        add_p(doc, "", indent=False)
    add_p(doc, "Лабораторная работа", align="center", indent=False, bold=True, size=Pt(16))
    add_p(doc, "Численные методы решения задач безусловной многопараметрической оптимизации", align="center", indent=False)
    add_p(doc, "По дисциплине «Методы оптимизации»", align="center", indent=False)
    add_p(doc, "Вариант 16", align="center", indent=False)
    for _ in range(5):
        add_p(doc, "", indent=False)
    add_p(doc, "Выполнил: студент гр. М25-787-1 Р. В. Скороходов", align="right", indent=False)
    add_p(doc, "Принял: доктор физико-математических наук, профессор В. А. Тененев", align="right", indent=False)
    for _ in range(3):
        add_p(doc, "", indent=False)
    add_p(doc, "Ижевск 2026", align="center", indent=False)
    page_break(doc)


def add_contents(doc):
    add_heading(doc, "Содержание")
    entries = [
        ("1", "Обозначения и ключевые понятия"),
        ("1.1", "Зачем эта работа"),
        ("2", "Цель работы"),
        ("3", "Постановка задачи"),
        ("4", "Аналитическое исследование функции"),
        ("5", "Численные методы оптимизации"),
        ("6", "Сравнительный анализ методов"),
        ("7", "Выводы"),
        ("", "Приложение А. Полные листинги реализаций"),
        ("", "Приложение Б. Вопросы и ответы для защиты"),
    ]
    for num, title in entries:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        text = f"{num}  {title}" if num else title
        run = p.add_run(clean_text(text))
        set_run(run)
        page = p.add_run("\t")
        set_run(page)
    page_break(doc)


def add_notation(doc):
    add_heading(doc, "1  Обозначения и ключевые понятия")
    add_table_caption(doc, "Таблица 1.1. Обозначения и ключевые понятия")
    rows = [
        ("x = (x, y)", "Точка на плоскости"),
        ("f(x, y)", "Значение функции в этой точке, то есть то, что минимизируем"),
        ("∇f", "Градиент функции, вектор направления самого быстрого роста"),
        ("H", "Матрица Гессе, характеризует кривизну функции в точке"),
        ("ε", "Заданная точность вычислений"),
        ("x_k", "Точка, полученная на k-й итерации"),
        ("x*", "Истинная точка минимума, то, что ищем"),
        ("f*", "Минимальное значение функции, f* = f(x*)"),
        ("Q, b", "Матрица и вектор в квадратичной форме функции"),
        ("d_k", "Направление движения метода на шаге k"),
        ("α_k", "Длина шага на k-й итерации"),
    ]
    add_table(doc, ["Символ", "Что означает"], rows, font_size=Pt(11))
    add_heading(doc, "1.1  Зачем эта работа", level=2)
    add_p(doc, "Оптимизация - поиск точки, где функция принимает наименьшее или наибольшее значение. К этой задаче сводятся подбор коэффициентов нейросетей, поиск кратчайшего маршрута, минимизация стоимости производства, расчет формы крыла самолета. В лабораторной берется простая функция двух переменных, для которой ответ можно найти на бумаге, и на ней проверяются четыре численных метода. Простая функция нужна не сама по себе, а как способ убедиться, что методы работают правильно, прежде чем применять их к сложным задачам, где ответ заранее неизвестен.")


def add_goal_and_problem(doc):
    add_heading(doc, "2  Цель работы")
    add_p(doc, "Цель работы - последовательно проверить численные методы безусловной оптимизации на функции двух переменных и сравнить их с аналитически найденным ответом. Я сначала нахожу минимум на бумаге, затем запускаю четыре метода из одной начальной точки и проверяю, совпадает ли результат.")

    add_heading(doc, "3  Постановка задачи")
    add_p(doc, "Согласно варианту 16, требуется найти безусловный минимум функции двух переменных. Формула ниже задает исходную функцию, именно ее минимум я дальше ищу аналитически и численно:")
    add_formula(doc, r"f(x,y)=2(x-3)^2+3(y+2)^2-2xy+4x-6y", "3.1")
    add_p(doc, "Функция определена на всей плоскости. Начальная точка для всех численных методов выбирается одинаковой: x0 = (0; 0). Точность вычислений принимается ε = 10^-4.")
    add_dash(doc, "построить поверхность функции и линии уровня;")
    add_dash(doc, "найти стационарную точку аналитически;")
    add_dash(doc, "проверить характер этой точки через матрицу Гессе;")
    add_dash(doc, "решить задачу четырьмя численными методами;")
    add_dash(doc, "сравнить траектории и число итераций.")


def add_analysis(doc):
    add_heading(doc, "4  Аналитическое исследование функции")
    add_heading(doc, "4.1  Раскрытие функции и градиент", level=2)
    add_p(doc, "Раскрытый вид функции удобнее для дифференцирования, потому что сразу видны квадратичные, смешанные и линейные слагаемые:")
    add_formula(doc, r"f(x,y)=2x^2+3y^2-2xy-8x+6y+30", "4.1")
    add_p(doc, "Частная производная по x показывает скорость изменения функции при движении вдоль оси x:")
    add_formula(doc, r"\frac{\partial f}{\partial x}=4x-2y-8", "4.2")
    add_p(doc, "Частная производная по y нужна вместе с предыдущей производной, чтобы составить систему ∇f = 0:")
    add_formula(doc, r"\frac{\partial f}{\partial y}=-2x+6y+6", "4.3")
    add_p(doc, "В основном тексте оставлена сигнатура функции градиента. Полная реализация приведена в Приложении А.")
    add_code_block(doc, '''def grad(point: np.ndarray | list[float] | tuple[float, float]) -> np.ndarray:
    """Return gradient of f at point."""
    # полная реализация приведена в Приложении А и в репозитории''')

    add_heading(doc, "4.2  Стационарная точка", level=2)
    add_p(doc, "Стационарная точка находится из условия ∇f = 0. Это точка, где функция перестает возрастать или убывать по отдельным координатам:")
    add_formula(doc, r"\left\{\substack{4x-2y-8=0\\-2x+6y+6=0}\right.")
    add_p(doc, "Решение системы дает единственную стационарную точку. Ее значение функции нужно как эталон для всех численных методов:")
    add_formula(doc, r"x^*=(1{,}8;\,-0{,}4),\quad f(x^*)=21{,}6", "4.4")

    add_heading(doc, "4.3  Матрица Гессе и проверка минимума", level=2)
    add_p(doc, "Вторые производные нужны, чтобы собрать матрицу Гессе и проверить характер стационарной точки:")
    add_formula(doc, r"\frac{\partial^2 f}{\partial x^2}=4,\quad \frac{\partial^2 f}{\partial y^2}=6,\quad \frac{\partial^2 f}{\partial x\,\partial y}=-2")
    add_p(doc, "Матрица Гессе показывает кривизну функции. Для данной функции она постоянна:")
    add_formula(doc, r"H=\left[\substack{4\quad -2\\-2\quad 6}\right]", "4.5")
    add_p(doc, "Положительные главные миноры подтверждают, что функция выпуклая и найденная точка является минимумом:")
    add_formula(doc, r"\Delta_1=4>0,\quad \Delta_2=\det H=4\cdot6-(-2)\cdot(-2)=20>0")
    add_p(doc, "Оба угловых минора положительны, значит H положительно определена. Следовательно, x* = (1,8; -0,4) - глобальный минимум. Точку x* мы нашли аналитически, на бумаге. В разделе 5 проверим, могут ли численные методы найти эту же точку, стартуя из x0 = (0; 0).")

    add_heading(doc, "4.4  График функции", level=2)
    add_figure(doc, "figures/fig_3_4_surface_contour.png", "Рисунок 4.1. Поверхность функции слева и ее линии уровня справа. Линии уровня - концентрические эллипсы; красная звезда отмечает точку минимума x* = (1,8; -0,4)")


def add_method_intro(doc, title, analogy, scheme):
    add_heading(doc, title, level=2)
    add_p(doc, analogy, italic=True)
    add_code_block(doc, scheme, size=Pt(10.5))


def add_remember(doc, items):
    add_p(doc, "Что нужно запомнить:", bold=True)
    for item in items:
        add_dash(doc, item)


def add_methods(doc, values):
    add_heading(doc, "5  Численные методы оптимизации")
    add_p(doc, "Прежде чем разбирать каждый метод по отдельности, посмотрим на общую картину: где какой метод применим и насколько быстро работает на нашей квадратичной функции. Эта таблица повторится в сравнительном анализе уже с более детальными цифрами.")
    add_table_caption(doc, "Таблица 5.1. Сравнение методов по применимости")
    add_table(
        doc,
        ["Метод", "Когда выбирать", "Нужны производные", "Скорость на нашей задаче"],
        [
            ["Нелдер-Мид", "Функция может быть негладкой или с разрывами", "Нет", f"Медленно ({values['nm_iter']} итерация)"],
            ["Хук-Дживс", "Производных нет, но функция достаточно гладкая", "Нет", f"Средне ({values['hj_iter']} шагов)"],
            ["Наискорейший спуск", "Производные доступны, функция простая", "Да, градиент", f"Средне ({values['sd_iter']} итераций)"],
            ["Флетчер-Ривс", "Функция квадратичная или близка к ней", "Да, градиент", f"Очень быстро ({values['fr_iter']} итерации)"],
        ],
        font_size=Pt(9.5),
    )

    add_method_intro(
        doc,
        "5.1  Метод Нелдера-Мида",
        "Группа из трех человек ищет самую низкую точку в овраге. На каждом шаге тот, кто стоит выше всех, перепрыгивает через центр двух остальных в надежде попасть ниже.",
        f"[x0=(0;0), eps=1e-4, step=1] -> [Нелдер-Мид] -> [x≈({values['nm_x']}; {values['nm_y']}), f≈{values['nm_f']}, итераций: {values['nm_iter']}]",
    )
    add_p(doc, "Идея метода: вместо производных используется треугольник из трех точек. Метод двигает этот треугольник так, чтобы его лучшая вершина постепенно приближалась к минимуму. Коэффициенты: отражение α = 1, растяжение γ = 2, сжатие β = 0,5, редукция σ = 0,5.")
    add_p(doc, "Данные для запуска: начальная точка, размер начального симплекса, точность ε и коэффициенты преобразования симплекса.")
    add_p(doc, "Одна итерация: вершины сортируются по значению функции, худшая вершина заменяется отраженной, растянутой или сжатой точкой.")
    add_p(doc, "Критерий остановки: разброс значений функции в вершинах симплекса становится меньше ε.")
    add_code_block(doc, signature_block("nelder_mead"))
    headers, rows = read_csv(ROOT / "tables" / "tab_4_1_nm.csv")
    add_table_caption(doc, "Таблица 5.2. Сходимость метода Нелдера-Мида")
    add_table(doc, headers, rows, font_size=Pt(9.2))
    add_p(doc, "Из таблицы видно, что значение функции постепенно уменьшается и приближается к 21,6. Координаты лучшей точки также приближаются к аналитически найденному минимуму x* = (1,8; -0,4). Разброс значений функции в вершинах симплекса становится меньше заданной точности, поэтому метод останавливается.")
    add_figure(doc, "figures/fig_4_1_nm_trajectory.png", "Рисунок 5.1. Треугольные симплексы метода Нелдера-Мида на линиях уровня. Симплекс постепенно схлопывается к красной звезде - найденному минимуму")
    add_p(doc, f"Что получилось в данной задаче: метод сошелся к точке ({values['nm_x']}; {values['nm_y']}) с f = {values['nm_f']} за {values['nm_iter']} итерацию. На графике видно, что симплекс перемещается к центру эллипсов линий уровня.")
    add_remember(doc, ["не требует производных функции;", "подходит для негладких и зашумленных функций;", "на гладких квадратичных задачах проигрывает градиентным методам в скорости."])

    add_method_intro(
        doc,
        "5.2  Метод Хука-Дживса",
        "Стоишь на склоне, смотришь на шаг в каждую сторону. Куда уклон вниз - туда идешь, пока спускается, а если улучшения нет - уменьшаешь шаг.",
        f"[x0=(0;0), eps=1e-4, delta=1] -> [Хук-Дживс] -> [x≈({values['hj_x']}; {values['hj_y']}), f≈{values['hj_f']}, шагов: {values['hj_iter']}]",
    )
    add_p(doc, "Идея метода: сначала проверяются небольшие шаги по координатам, затем при успешном движении выполняется шаг по найденному направлению.")
    add_p(doc, "Данные для запуска: начальная точка, координатный шаг Δ, точность ε и метод одномерного поиска.")
    add_p(doc, "Одна итерация: метод пробует сдвиги +Δ и -Δ по каждой координате, после улучшения выбирает направление и уточняет длину шага методом деления отрезка пополам.")
    add_p(doc, "Критерий остановки: координатный шаг Δ становится меньше ε.")
    add_code_block(doc, signature_block("bisection_min"))
    add_code_block(doc, signature_block("hooke_jeeves"))
    headers, rows = read_csv(ROOT / "tables" / "tab_4_2_hj.csv")
    add_table_caption(doc, "Таблица 5.3. Сходимость метода Хука-Дживса")
    add_table(doc, headers, rows, font_size=Pt(9.2))
    add_p(doc, "Из таблицы видно, что координаты базовой точки постепенно уточняются, а значение функции становится практически равным 21,6. Уменьшение Δ показывает, что метод перестал находить заметные улучшения крупными координатными шагами. Это и является основанием для остановки.")
    add_figure(doc, "figures/fig_4_2_hj_trajectory.png", "Рисунок 5.2. Базовые точки метода Хука-Дживса, синие кружки соединены линией на линиях уровня. От стартовой точки (0; 0) идем к минимуму x* за 25 шагов")
    add_p(doc, f"Что получилось в данной задаче: метод сошелся к точке ({values['hj_x']}; {values['hj_y']}) с f = {values['hj_f']} за {values['hj_iter']} шагов. На графике видно последовательное приближение базовых точек к центру линий уровня.")
    add_remember(doc, ["производные тоже не нужны, как и у Нелдера-Мида;", "использует одномерный поиск, у нас - деление отрезка пополам;", "более эффективен, чем Нелдер-Мид, на гладких функциях."])

    add_method_intro(
        doc,
        "5.3  Метод наискорейшего спуска",
        "Шарик катится туда, где склон самый крутой в данной точке.",
        f"[x0=(0;0), eps=1e-4, grad, Q] -> [Наискорейший спуск] -> [x≈({values['sd_x']}; {values['sd_y']}), f≈{values['sd_f']}, итераций: {values['sd_iter']}]",
    )
    add_p(doc, "Идея метода: в каждой точке идти против градиента, потому что градиент показывает направление наибольшего роста функции.")
    add_p(doc, "Данные для запуска: начальная точка, градиент функции, матрица Q для вычисления оптимального шага и точность ε.")
    add_p(doc, "Одна итерация: вычисляется градиент, выбирается шаг α_k по расчетной формуле, затем строится новая точка x_{k+1}.")
    add_p(doc, "Критерий остановки: норма градиента становится меньше ε.")
    add_p(doc, "Метод применим, так как f квадратичная, это показано в разделе 4.")
    add_p(doc, "Квадратичная запись нужна для градиентных методов, потому что шаг можно вычислять через матрицу Q:")
    add_formula(doc, r"f(x)=\frac{1}{2}x^TQx+b^Tx+c", "5.1")
    add_p(doc, "Значения Q, b и c связывают исходную функцию с формулами градиентных методов:")
    add_formula(doc, r"Q=\left[\substack{4\quad -2\\-2\quad 6}\right],\quad b=\left(\substack{-8\\6}\right),\quad c=30")
    add_p(doc, "Градиент в компактной форме:")
    add_formula(doc, r"\nabla f(x)=Qx+b", "5.2")
    add_p(doc, "Формула показывает, как метод получает новую точку и как выбирает оптимальный шаг:")
    add_formula(doc, r"x_{k+1}=x_k-\alpha_k\nabla f(x_k),\quad \alpha_k=\frac{g^Tg}{g^TQg}", "5.3")
    add_code_block(doc, signature_block("steepest_descent"))
    headers, rows = read_csv(ROOT / "tables" / "tab_4_3_sd.csv")
    add_table_caption(doc, "Таблица 5.4. Сходимость метода наискорейшего спуска")
    add_table(doc, headers, rows, font_size=Pt(9.2))
    add_p(doc, "Из таблицы видно, что значение нормы градиента уменьшается, значит точка становится ближе к стационарной. Значение функции приближается к 21,6, а координаты стремятся к (1,8; -0,4). Метод останавливается тогда, когда градиент становится достаточно малым.")
    add_figure(doc, "figures/fig_4_3_sd_trajectory.png", "Рисунок 5.3. Траектория метода наискорейшего спуска. Видно характерное зигзагообразное движение: каждый отрезок траектории почти перпендикулярен предыдущему")
    add_p(doc, f"Что получилось в данной задаче: метод сошелся за {values['sd_iter']} итераций к точке ({values['sd_x']}; {values['sd_y']}) с f = {values['sd_f']}. На графике видно зигзагообразное движение к минимуму.")
    add_remember(doc, ["использует градиент: нужна аналитическая производная;", "для квадратичной функции длина шага находится явно;", "дает характерную зигзагообразную траекторию из-за ортогональности соседних шагов."])

    add_method_intro(
        doc,
        "5.4  Метод Флетчера-Ривса",
        "Тот же шарик, но запоминает, куда катился раньше, и не разворачивается на 90 градусов.",
        f"[x0=(0;0), eps=1e-4, grad, Q] -> [Флетчер-Ривс] -> [x=({values['fr_x']}; {values['fr_y']}), f={values['fr_f']}, итераций: {values['fr_iter']}]",
    )
    add_p(doc, "Идея метода: двигаться не просто против градиента, а по сопряженным направлениям, которые учитывают кривизну квадратичной функции.")
    add_p(doc, "Данные для запуска: начальная точка, градиент, матрица Q, точность ε и формула для коэффициента β_k.")
    add_p(doc, "Одна итерация: вычисляется шаг вдоль текущего направления, затем строится новое направление из антиградиента и части предыдущего направления.")
    add_p(doc, "Критерий остановки: норма градиента становится меньше ε.")
    add_p(doc, "Первое направление берется против градиента, то есть в сторону убывания функции:")
    add_formula(doc, r"d_0=-\nabla f(x_0)", "5.4")
    add_p(doc, "По этой формуле находится длина шага вдоль текущего направления:")
    add_formula(doc, r"\alpha_k=-\frac{g^Td_k}{d_k^TQd_k}", "5.5")
    add_p(doc, "Так метод переходит из текущей точки в следующую:")
    add_formula(doc, r"x_{k+1}=x_k+\alpha_k d_k", "5.6")
    add_p(doc, "Коэффициент β_k показывает, какую часть предыдущего направления нужно сохранить:")
    add_formula(doc, r"\beta_k=\frac{\Vert g_{k+1}\Vert^2}{\Vert g_k\Vert^2}", "5.7")
    add_p(doc, "Новое направление строится с учетом предыдущего направления, поэтому метод быстрее для квадратичных функций:")
    add_formula(doc, r"d_{k+1}=-g_{k+1}+\beta_kd_k", "5.8")
    add_code_block(doc, signature_block("fletcher_reeves"))
    headers, rows = read_csv(ROOT / "tables" / "tab_4_4_fr.csv")
    add_table_caption(doc, "Таблица 5.5. Сходимость метода Флетчера-Ривса")
    add_table(doc, headers, rows, font_size=Pt(9.2))
    add_p(doc, "Из таблицы видно, что метод Флетчера-Ривса быстро уменьшает норму градиента и уже за две итерации попадает в точку минимума. Для этой задачи такой результат ожидаем, потому что функция квадратичная и имеет две переменные.")
    add_figure(doc, "figures/fig_4_4_fr_trajectory.png", "Рисунок 5.4. Траектория Флетчера-Ривса: от (0; 0) две прямые - и сразу в минимум. Это квадратичное завершение в действии")
    add_p(doc, f"Что получилось в данной задаче: метод сошелся к точному решению ({values['fr_x']}; {values['fr_y']}) с f = {values['fr_f']} за {values['fr_iter']} итерации. Все четыре метода сошлись к (1,8; -0,4), но за разное число шагов. В разделе 6 сравним, кто оказался эффективнее.")
    add_remember(doc, ["метод сопряженных градиентов накапливает информацию о предыдущих направлениях;", "для квадратичной функции n переменных сходится не более чем за n шагов;", "в нашей задаче n = 2, поэтому 2 итерации - это закономерный результат."])


def add_comparison_and_conclusion(doc):
    add_heading(doc, "6  Сравнительный анализ методов")
    headers, rows = read_csv(ROOT / "tables" / "tab_5_compare.csv")
    add_table_caption(doc, "Таблица 6.1. Сравнение численных методов оптимизации")
    add_table(doc, headers, rows, font_size=Pt(9.2))
    add_p(doc, "Из таблицы видно, что все методы приходят к одному и тому же минимуму, но делают это с разным числом итераций. Самым быстрым оказался метод Флетчера-Ривса, а метод Нелдера-Мида потребовал больше шагов. Поэтому таблица сравнивает не только правильность ответа, но и эффективность способов поиска.")
    add_figure(doc, "figures/fig_5_convergence.png", "Рисунок 6.1. Сходимость всех четырех методов: ось Y в логарифмическом масштабе показывает f(x_k) - f*, ось X - номер итерации. Чем круче падает кривая, тем быстрее метод")
    add_p(doc, "Теперь, когда все методы испытаны на одной задаче, можно сформулировать общие закономерности.")

    add_heading(doc, "7  Выводы")
    add_p(doc, "В ходе работы я не просто применил четыре метода, а последовательно проверил один и тот же результат разными способами: аналитическим расчетом, таблицами итераций и графиками траекторий. Совпадение этих трех проверок показывает, что найденный минимум не является случайным результатом одного алгоритма.")
    add_dash(doc, "аналитически найден глобальный минимум x* = (1,8; -0,4), f* = 21,6;")
    add_dash(doc, "все четыре численных метода пришли к этой же точке с заданной точностью;")
    add_dash(doc, "метод Флетчера-Ривса оказался самым быстрым, потому что функция квадратичная и имеет две переменные;")
    add_dash(doc, "методы без производных полезны, когда градиент неизвестен или функция негладкая;")
    add_dash(doc, "графики подтверждают характер движения каждого метода, а таблицы объясняют момент остановки.")


def add_appendix_a(doc):
    page_break(doc)
    add_heading(doc, "Приложение А. Полные листинги реализаций")
    add_p(doc, "В основном тексте оставлены только сигнатуры функций. Ниже приведены полные реализации пяти функций, использованных в ноутбуках и при построении таблиц.")
    functions = [
        ("src/nelder_mead.py", "nelder_mead"),
        ("src/hooke_jeeves.py", "bisection_min"),
        ("src/hooke_jeeves.py", "hooke_jeeves"),
        ("src/gradient_methods.py", "steepest_descent"),
        ("src/gradient_methods.py", "fletcher_reeves"),
    ]
    for file_name, func_name in functions:
        add_p(doc, f"Функция {func_name} из файла {file_name}", bold=True)
        add_code_block(doc, extract_function_source(ROOT / file_name, func_name), size=Pt(7.3))


def add_appendix_b(doc):
    add_heading(doc, "Приложение Б. Вопросы и ответы для защиты")
    add_p(doc, "Здесь собраны вопросы, которые чаще всего задают по этой работе, и краткие ответы на них.")
    qa = [
        ("Вопрос 1. Что такое стационарная точка?", "Точка, где градиент функции равен нулю: ∇f = 0. Это либо локальный минимум, либо максимум, либо седловая точка. Тип точки определяется матрицей Гессе."),
        ("Вопрос 2. Зачем нужна положительная определенность гессиана?", "Если все собственные значения матрицы Гессе положительны, функция в этой точке имеет локальный минимум. Если все отрицательны - максимум. Если знаки разные - седло."),
        ("Вопрос 3. Почему метод Флетчера-Ривса сошелся за 2 итерации?", "Это свойство квадратичного завершения. Для квадратичной функции n переменных метод сопряженных градиентов сходится не более чем за n шагов. У нашей функции n = 2, поэтому 2 итерации - это гарантия, а не везение."),
        ("Вопрос 4. Почему наискорейший спуск дает зигзаг?", "Длина шага выбирается так, что в новой точке градиент перпендикулярен старому направлению. Поэтому соседние ходы идут почти под прямым углом. На вытянутых эллипсах линий уровня это приводит к движению змейкой."),
        ("Вопрос 5. В чем преимущество Хука-Дживса перед Нелдером-Мидом?", "Хук-Дживс использует одномерный поиск для нахождения оптимальной длины шага по образцу. За счет этого он эффективнее на гладких функциях. Нелдер-Мид зато не использует одномерного поиска и устойчивее на негладких функциях."),
        ("Вопрос 6. Что значит метод деления отрезка пополам?", "Это способ найти минимум функции одной переменной на отрезке [a, b]. На каждом шаге берутся середина и две точки в четвертях отрезка. После сравнения значений функции отбрасывается часть отрезка, где минимума нет."),
        ("Вопрос 7. Почему используется точность 10^-4, а не 10^-8?", "Точности 10^-4 достаточно, чтобы все методы сошлись к точному ответу с погрешностью меньше 0,001. Большая точность увеличит число итераций, но не даст новой содержательной информации для этой лабораторной."),
        ("Вопрос 8. Что произойдет, если стартовать из другой точки?", "Поскольку функция выпуклая, все методы должны прийти к тому же x* = (1,8; -0,4). Изменится только число итераций. Это легко проверить, поменяв x0 в ноутбуке."),
    ]
    for question, answer in qa:
        add_p(doc, question, bold=True)
        add_p(doc, answer)


def all_doc_text(doc: Document) -> str:
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(p.text for p in cell.paragraphs)
    return "\n".join(texts)


def run_acceptance(doc: Document) -> None:
    text = all_doc_text(doc)
    assert chr(0x2014) not in text, "Найдены em-dash"
    assert chr(0x2013) not in text, "Найдены en-dash"
    assert not re.search(r"^Зачем нужна формула", text, flags=re.MULTILINE)
    expected_sections = [
        "Обозначения и ключевые понятия",
        "Цель работы",
        "Постановка задачи",
        "Аналитическое исследование функции",
        "Численные методы оптимизации",
        "Сравнительный анализ методов",
        "Выводы",
        "Приложение А",
        "Приложение Б",
    ]
    for section in expected_sections:
        assert any(section in p.text for p in doc.paragraphs), f"Не найден раздел: {section}"
    for marker in ["перепрыгивает через центр", "идешь, пока спускается", "склон самый крутой", "не разворачивается"]:
        assert any(marker in p.text for p in doc.paragraphs), f"Не найдена аналогия: {marker}"
    remember_count = sum(1 for p in doc.paragraphs if "Что нужно запомнить" in p.text)
    assert remember_count == 4, f"Ожидалось 4 блока, найдено {remember_count}"
    qa_count = sum(1 for p in doc.paragraphs if p.text.startswith("Вопрос "))
    assert qa_count == 8, f"Ожидалось 8 вопросов, найдено {qa_count}"
    appendix_start = max(i for i, p in enumerate(doc.paragraphs) if p.text.startswith("Приложение А"))
    main_paragraphs = doc.paragraphs[:appendix_start]
    for func_name in ["nelder_mead", "bisection_min", "hooke_jeeves", "steepest_descent", "fletcher_reeves"]:
        main_blocks = [p.text for p in main_paragraphs if f"def {func_name}" in p.text]
        assert main_blocks, f"Нет сигнатуры {func_name}"
        assert all(len(block.splitlines()) <= 8 for block in main_blocks), f"Длинный основной блок {func_name}"
        assert any(f"def {func_name}" in p.text for p in doc.paragraphs[appendix_start:]), f"Нет полного листинга {func_name}"


def build() -> None:
    values = load_values()
    doc = Document()
    setup_document(doc)
    add_title(doc)
    add_contents(doc)
    add_notation(doc)
    add_goal_and_problem(doc)
    add_analysis(doc)
    add_methods(doc, values)
    add_comparison_and_conclusion(doc)
    add_appendix_a(doc)
    add_appendix_b(doc)
    run_acceptance(doc)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print("[Acceptance] no long dashes: ok")
    print("[Acceptance] expected sections: ok")
    print("[Acceptance] analogies: ok")
    print("[Acceptance] remember blocks: ok")
    print("[Acceptance] FAQ: ok")


if __name__ == "__main__":
    build()
