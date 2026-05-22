"""Сборка шаблона отчёта lab1_v16. ГОСТ 7.32-2017 + ИжГТУ титульник.

Метки-плейсхолдеры (build_report.py заменит их по факту):
  [[VAL:key]]                      — инлайн-значение из values.json
  [[FIG:path|caption]]             — рисунок + подпись
  [[TBL:path|caption]]             — таблица + подпись (caption сверху по ГОСТ)
  [[CODE:src/<file>.py|<name>]]   — листинг функции из исходника
"""
import hashlib
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor

FONT = "Times New Roman"
SZ = Pt(14)
SZ_SMALL = Pt(12)
ROOT = Path(__file__).parent
FORMULA_DIR = ROOT / "formula_images"


def set_run(run, text, *, bold=False, italic=False, size=SZ, font=FONT, mono=False):
    run.text = text
    run.font.name = "Consolas" if mono else font
    run.font.size = size
    run.bold = bold
    run.italic = italic
    # CJK тоже Times New Roman иначе кириллица иногда едет
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), "Consolas" if mono else font)
    rFonts.set(qn('w:hAnsi'), "Consolas" if mono else font)
    rFonts.set(qn('w:ascii'), "Consolas" if mono else font)


def add_p(doc, text="", *, align=None, indent_first=True, line_spacing=1.5,
          space_after=0, bold=False, italic=False, mono=False, size=SZ):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if indent_first and not align:
        pf.first_line_indent = Cm(1.25)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        run = p.add_run()
        set_run(run, text, bold=bold, italic=italic, mono=mono, size=size)
    return p


def add_heading(doc, text, *, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run()
    set_run(run, text, bold=True, size=SZ)
    return p


def add_dash(doc, text):
    """Перечисление через дефис по ГОСТ 7.32-2017."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-0.6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run()
    set_run(run, "- " + text)
    return p


def add_placeholder(doc, text):
    """Видимая метка-плейсхолдер для скрипта-наполнителя."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run(run, text, italic=True, mono=True, size=Pt(11))
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def render_formula(latex, *, fontsize=18):
    """Render a mathtext/LaTeX-like formula to a transparent PNG."""
    FORMULA_DIR.mkdir(exist_ok=True)
    digest = hashlib.sha1(latex.encode("utf-8")).hexdigest()[:16]
    path = FORMULA_DIR / f"formula_{digest}.png"
    if path.exists():
        return path

    # Matplotlib mathtext does not require an external TeX installation and is
    # stable on clean Windows machines. Matrix-like expressions use \substack.
    fig = plt.figure(figsize=(9.0, 1.15), dpi=300)
    fig.text(0.5, 0.5, f"${latex}$", ha="center", va="center", fontsize=fontsize)
    fig.savefig(path, transparent=True, bbox_inches="tight", pad_inches=0.04)
    plt.close(fig)
    return path


def add_formula(doc, latex, *, number=None):
    """Формула как LaTeX-rendered PNG; при наличии номер ставится справа."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    formula_path = render_formula(latex)
    if number:
        # Левый отступ нулевой, формула по центру, номер справа.
        pf.tab_stops.add_tab_stop(Cm(8.0), WD_TAB_ALIGNMENT.CENTER)
        pf.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("\t")
        run = p.add_run()
        run.add_picture(str(formula_path))
        run2 = p.add_run(f"\t({number})")
        set_run(run2, "\t" + f"({number})")
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(formula_path))
    return p


def setup_section(doc):
    s = doc.sections[0]
    s.orientation = WD_ORIENTATION.PORTRAIT
    s.page_width = Mm(210)
    s.page_height = Mm(297)
    s.top_margin = Mm(20)
    s.bottom_margin = Mm(20)
    s.left_margin = Mm(30)
    s.right_margin = Mm(15)


def add_page_number(doc):
    """Номер страницы по центру нижнего поля."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    # Поле PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run(run, "")
    run.text = ""


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


# ===========================================================
doc = Document()
setup_section(doc)

# Стиль по умолчанию
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = SZ

# Номера страниц
add_page_number(doc)

# ----------- ТИТУЛЬНЫЙ ЛИСТ -----------
add_p(doc, "Министерство образования и науки Российской Федерации", align="center", indent_first=False)
add_p(doc, "Федеральное государственное бюджетное образовательное учреждение высшего образования", align="center", indent_first=False)
add_p(doc, "«Ижевский государственный технический университет имени М.Т. Калашникова»", align="center", indent_first=False)
for _ in range(6):
    add_p(doc, "", indent_first=False)
add_p(doc, "Лабораторная работа", align="center", indent_first=False, bold=True, size=Pt(16))
add_p(doc, "", indent_first=False)
add_p(doc, "Численные методы решения задач безусловной многопараметрической оптимизации", align="center", indent_first=False)
add_p(doc, "", indent_first=False)
add_p(doc, "По дисциплине «Методы оптимизации»", align="center", indent_first=False)
add_p(doc, "", indent_first=False)
add_p(doc, "Вариант № 16", align="center", indent_first=False)
for _ in range(6):
    add_p(doc, "", indent_first=False)

# Подписи через табы
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
run = p.add_run("Выполнил:")
set_run(run, "Выполнил:")
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
run = p.add_run("Студент гр. М25-787-1\t")
set_run(run, "Студент гр. М25-787-1\t")
run = p.add_run("Р. В. Скороходов")
set_run(run, "Р. В. Скороходов")
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run("Принял:")
set_run(run, "Принял:")
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
tab_stops = p.paragraph_format.tab_stops
tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT)
run = p.add_run("доктор физико-математических наук, профессор\t")
set_run(run, "доктор физико-математических наук, профессор\t")
run = p.add_run("В. А. Тененев")
set_run(run, "В. А. Тененев")
add_p(doc, "", indent_first=False)
add_p(doc, "Ижевск 2026", align="center", indent_first=False)

# ------------- СОДЕРЖАНИЕ ---------------
page_break(doc)
add_heading(doc, "Содержание")
toc_entries = [
    ("1", "Цель работы"),
    ("2", "Постановка задачи"),
    ("3", "Аналитическое исследование функции"),
    ("3.1", "Раскрытие функции и градиент"),
    ("3.2", "Стационарные точки"),
    ("3.3", "Гессиан и характер экстремума"),
    ("3.4", "Графики поверхности и линий уровня"),
    ("4", "Численные методы оптимизации"),
    ("4.1", "Метод Нелдера–Мида"),
    ("4.2", "Метод Хука–Дживса"),
    ("4.3", "Метод наискорейшего спуска"),
    ("4.4", "Метод Флетчера–Ривса"),
    ("5", "Сравнительный анализ методов"),
    ("6", "Выводы"),
    ("",  "Приложение А. Репозиторий с исходным кодом и ноутбуками"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    text = f"{num}  {title}" if num else title
    run = p.add_run(text)
    set_run(run, text)
    run = p.add_run("\t[[VAL:page_" + (num.replace(".", "_") or "app") + "]]")
    set_run(run, "\t[[VAL:page_" + (num.replace(".", "_") or "app") + "]]", size=Pt(11), italic=True)

# ============= 1. ЦЕЛЬ ===============
page_break(doc)
add_heading(doc, "1  Цель работы")
add_p(doc, "Изучить численные методы решения задач безусловной многопараметрической "
      "оптимизации и сравнить их по числу итераций и точности на квадратичной функции "
      "двух переменных. В работе реализуются и сопоставляются четыре метода: "
      "Нелдера–Мида, Хука–Дживса с одномерной оптимизацией шага методом деления отрезка "
      "пополам, метод наискорейшего спуска и метод сопряжённых градиентов Флетчера–Ривса. "
      "Все вычисления и графики выполняются в Python-ноутбуках, опубликованных в репозитории "
      "на GitHub; каждый ноутбук открывается одним кликом в Google Colab и выполняется без "
      "локальной установки.")

# ============= 2. ПОСТАНОВКА ===============
add_heading(doc, "2  Постановка задачи")
add_p(doc, "Согласно варианту 16, требуется найти безусловный минимум функции двух переменных")
add_formula(doc, r"f(x,y)=2(x-3)^2+3(y+2)^2-2xy+4x-6y,", number="2.1")
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
run = p.add_run("определённой на всей плоскости ℝ². В работе необходимо:")
set_run(run, "определённой на всей плоскости ℝ². В работе необходимо:")
add_dash(doc, "построить график функции и линии уровня;")
add_dash(doc, "найти стационарные точки аналитически и классифицировать их;")
add_dash(doc, "решить задачу безусловной оптимизации методом Нелдера–Мида с коэффициентом отражения α = 1 (остальные коэффициенты выбираются самостоятельно);")
add_dash(doc, "решить методом Хука–Дживса с одномерной оптимизацией шага методом деления отрезка пополам;")
add_dash(doc, "решить методом наискорейшего спуска и методом Флетчера–Ривса с явными расчётными формулами шага;")
add_dash(doc, "для каждого метода построить траекторию поиска на плоскости.")
add_p(doc, "Принимаем начальную точку x₀ = (0; 0) и точность ε = 10⁻⁴. Все вычисления "
      "выполнены в Python (NumPy + SymPy + Matplotlib); исходный код и ноутбуки доступны "
      "в репозитории GitHub (см. приложение А).")

# ============= 3. АНАЛИТИКА ===============
add_heading(doc, "3  Аналитическое исследование функции")

add_heading(doc, "3.1  Раскрытие функции и градиент", level=2)
add_p(doc, "Раскроем квадраты в выражении (2.1):")
add_formula(doc, r"f(x,y)=2x^2+3y^2-2xy-8x+6y+30.", number="3.1")
add_p(doc, "Частные производные первого порядка вычисляются через SymPy "
      "(см. ноутбук 01_analysis.ipynb):")
add_placeholder(doc, "[[CODE:src/target.py|grad]]")
add_p(doc, "В развёрнутом виде:")
add_formula(doc, r"\frac{\partial f}{\partial x}=4x-2y-8,", number="3.2")
add_formula(doc, r"\frac{\partial f}{\partial y}=-2x+6y+6.", number="3.3")

add_heading(doc, "3.2  Стационарные точки", level=2)
add_p(doc, "Координаты стационарных точек находятся из условия ∇f = 0, то есть из системы:")
add_formula(doc, r"\left\{\substack{4x-2y-8=0,\\-2x+6y+6=0.}\right.")
add_p(doc, "Из первого уравнения y = 2x − 4. Подстановка во второе даёт 10x = 18, "
      "откуда x = 9/5 = 1,8 и y = 2·1,8 − 4 = −0,4. Единственная стационарная точка:")
add_formula(doc, r"x^*=(1{,}8;\,-0{,}4),\quad f(x^*)=21{,}6.", number="3.4")

add_heading(doc, "3.3  Гессиан и характер экстремума", level=2)
add_p(doc, "Вторые частные производные постоянны (функция квадратичная):")
add_formula(doc, r"\frac{\partial^2 f}{\partial x^2}=4,\quad \frac{\partial^2 f}{\partial y^2}=6,\quad \frac{\partial^2 f}{\partial x\,\partial y}=-2.")
add_p(doc, "Матрица Гессе:")
add_formula(doc, r"H=\left[\substack{4\quad -2\\-2\quad 6}\right].", number="3.5")
add_p(doc, "По критерию Сильвестра проверим положительную определённость:")
add_formula(doc, r"\Delta_1=4>0,\quad \Delta_2=\det H=4\cdot 6-(-2)\cdot(-2)=20>0.")
add_p(doc, "Оба угловых минора положительны, значит H положительно определена. "
      "Следовательно, точка x* — точка локального минимума. Так как f(x) — выпуклая "
      "квадратичная (матрица квадратичной формы совпадает с H и положительно определена), "
      "x* является глобальным минимумом.")

add_heading(doc, "3.4  Графики поверхности и линий уровня", level=2)
add_p(doc, "Поверхность и линии уровня построены в matplotlib (ноутбук 01_analysis.ipynb):")
add_placeholder(doc, "[[FIG:figures/fig_3_4_surface_contour.png|Поверхность f(x, y) и её линии уровня; красной звездой обозначена точка минимума x* = (1,8;  −0,4)]]")
add_p(doc, "Линии уровня — концентрические эллипсы с центром в точке x*. Их вытянутость "
      "по диагонали отражает ненулевой коэффициент при xy в функции.")

# ============= 4. МЕТОДЫ ===============
add_heading(doc, "4  Численные методы оптимизации")
add_p(doc, "Во всех методах принимается начальная точка x₀ = (0; 0) и точность ε = 10⁻⁴. "
      "Истинный минимум функции — точка x* = (1,8; −0,4) с f* = 21,6 (раздел 3.2). "
      "Реализации алгоритмов лежат в модулях `src/`, ноутбуки в каталоге `notebooks/` "
      "репозитория содержат прогон каждого метода с записью истории итераций и построением "
      "траектории.")

# 4.1 Нелдер-Мид
add_heading(doc, "4.1  Метод Нелдера–Мида", level=2)
add_p(doc, "Метод деформируемого симплекса. Для функции n переменных строится симплекс "
      "из (n+1) вершины; в двумерном случае — треугольник. Используются коэффициенты:")
add_dash(doc, "α = 1: коэффициент отражения (по условию задания);")
add_dash(doc, "γ = 2: коэффициент растяжения;")
add_dash(doc, "β = 0,5: коэффициент сжатия;")
add_dash(doc, "σ = 0,5: коэффициент редукции.")
add_p(doc, "На каждой итерации: вершины сортируются по возрастанию f, худшая отражается "
      "через центр тяжести двух лучших. По значению f в отражённой точке выбирается "
      "растяжение, замена, сжатие или редукция. Критерий останова — разброс f-значений "
      "в вершинах симплекса меньше ε.")
add_p(doc, "Реализация на Python:")
add_placeholder(doc, "[[CODE:src/nelder_mead.py|nelder_mead]]")
add_placeholder(doc, "[[TBL:tables/tab_4_1_nm.csv|Сходимость метода Нелдера–Мида]]")
add_placeholder(doc, "[[FIG:figures/fig_4_1_nm_trajectory.png|Последовательность симплексов метода Нелдера–Мида на линиях уровня]]")
add_p(doc, "Метод сошёлся к точке x ≈ ([[VAL:nm_x]]; [[VAL:nm_y]]) с f ≈ [[VAL:nm_f]] "
      "за [[VAL:nm_iter]] итерации.")

# 4.2 Хук-Дживс
add_heading(doc, "4.2  Метод Хука–Дживса", level=2)
add_p(doc, "Метод не требует производных и сочетает исследующий поиск вокруг базовой точки "
      "с поиском по образцу вдоль найденного направления успеха. Длина шага по образцу "
      "определяется методом одномерной оптимизации, в данной задаче — методом деления "
      "отрезка пополам.")

add_heading(doc, "4.2.1  Метод деления отрезка пополам", level=2)
add_p(doc, "Минимизация унимодальной функции φ(t) на отрезке [a, b]: на каждой итерации "
      "вычисляются три точки x_m = (a+b)/2, x₁ = a + (b−a)/4, x₂ = b − (b−a)/4. "
      "Сравнением φ(x₁), φ(x_m), φ(x₂) отбрасывается одна треть отрезка. Останов "
      "при (b − a) < ε.")
add_placeholder(doc, "[[CODE:src/hooke_jeeves.py|bisection_min]]")

add_heading(doc, "4.2.2  Алгоритм Хука–Дживса", level=2)
add_p(doc, "Исследующий поиск: для каждой координаты пробуем приращения ±Δ и принимаем "
      "то, что уменьшает f. Если f уменьшилось — шаг по образцу в направлении "
      "d = x_new − x_base; оптимальная длина t* находится одномерной минимизацией "
      "φ(t) = f(x_new + t·d) методом деления отрезка пополам. Если улучшения не "
      "было — шаг Δ уменьшается вдвое.")
add_placeholder(doc, "[[CODE:src/hooke_jeeves.py|hooke_jeeves]]")
add_placeholder(doc, "[[TBL:tables/tab_4_2_hj.csv|Сходимость метода Хука–Дживса]]")
add_placeholder(doc, "[[FIG:figures/fig_4_2_hj_trajectory.png|Траектория базовых точек метода Хука–Дживса на линиях уровня]]")
add_p(doc, "Метод сошёлся к точке ([[VAL:hj_x]]; [[VAL:hj_y]]) с f = [[VAL:hj_f]] "
      "за [[VAL:hj_iter]] шагов по образцу.")

# 4.3 Наискорейший спуск
add_heading(doc, "4.3  Метод наискорейшего спуска", level=2)
add_p(doc, "Поскольку функция (2.1) квадратичная, её можно представить в виде")
add_formula(doc, r"f(x)=\frac{1}{2}x^TQx+b^Tx+c,", number="4.1")
add_p(doc, "В этой записи матрица квадратичной формы, вектор линейных коэффициентов и свободный член равны:")
add_formula(doc, r"Q=\left[\substack{4\quad -2\\-2\quad 6}\right],\quad b=\left(\substack{-8\\6}\right),\quad c=30.")
add_p(doc, "Градиент тогда")
add_formula(doc, r"\nabla f(x)=Qx+b.", number="4.2")
add_p(doc, "Итерация наискорейшего спуска и явная формула оптимального шага для квадратичной функции:")
add_formula(doc, r"x_{k+1}=x_k-\alpha_k\nabla f(x_k),\quad \alpha_k=\frac{g^Tg}{g^TQg}.", number="4.3")
add_p(doc, "Критерий останова: ‖∇f(x_k)‖ < ε.")
add_placeholder(doc, "[[CODE:src/gradient_methods.py|steepest_descent]]")
add_placeholder(doc, "[[TBL:tables/tab_4_3_sd.csv|Сходимость метода наискорейшего спуска]]")
add_placeholder(doc, "[[FIG:figures/fig_4_3_sd_trajectory.png|Траектория метода наискорейшего спуска; видна характерная зигзагообразная сходимость]]")
add_p(doc, "Метод сошёлся за [[VAL:sd_iter]] итераций к точке ([[VAL:sd_x]]; [[VAL:sd_y]]), "
      "f = [[VAL:sd_f]]. Зигзагообразная траектория — следствие ортогональности соседних "
      "направлений антиградиента и умеренной обусловленности матрицы Q (отношение собственных "
      "значений ≈ 2,6).")

# 4.4 Флетчер-Ривс
add_heading(doc, "4.4  Метод Флетчера–Ривса", level=2)
add_p(doc, "Метод сопряжённых градиентов. Направление спуска d_{k+1} строится как линейная "
      "комбинация антиградиента и предыдущего направления так, чтобы все направления были "
      "Q-сопряжёнными. Это даёт квадратичное завершение: для квадратичной функции n "
      "переменных метод сходится не более чем за n итераций. В данной задаче n = 2.")
add_p(doc, "Расчётные формулы:")
add_formula(doc, r"d_0=-\nabla f(x_0);", number="4.4")
add_formula(doc, r"\alpha_k=-\frac{g^Td_k}{d_k^TQd_k};", number="4.5")
add_formula(doc, r"x_{k+1}=x_k+\alpha_k d_k;", number="4.6")
add_formula(doc, r"\beta_k=\frac{\Vert g_{k+1}\Vert^2}{\Vert g_k\Vert^2};", number="4.7")
add_formula(doc, r"d_{k+1}=-g_{k+1}+\beta_k d_k.", number="4.8")
add_placeholder(doc, "[[CODE:src/gradient_methods.py|fletcher_reeves]]")
add_placeholder(doc, "[[TBL:tables/tab_4_4_fr.csv|Сходимость метода Флетчера–Ривса]]")
add_placeholder(doc, "[[FIG:figures/fig_4_4_fr_trajectory.png|Траектория метода Флетчера–Ривса; минимум достигается за 2 итерации (квадратичное завершение)]]")
add_p(doc, "Метод сошёлся к точному решению ([[VAL:fr_x]]; [[VAL:fr_y]]), f = [[VAL:fr_f]] "
      "ровно за [[VAL:fr_iter]] итерации, что подтверждает свойство квадратичного завершения.")

# ============= 5. СРАВНЕНИЕ ===============
add_heading(doc, "5  Сравнительный анализ методов")
add_p(doc, "В таблице 5.1 сведены результаты всех четырёх методов при общей начальной "
      "точке x₀ = (0; 0) и точности ε = 10⁻⁴.")
add_placeholder(doc, "[[TBL:tables/tab_5_compare.csv|Сравнение численных методов оптимизации]]")
add_placeholder(doc, "[[FIG:figures/fig_5_convergence.png|Зависимость лучшего значения f от номера итерации для всех методов (полу-логарифмический масштаб по f − f*)]]")

# ============= 6. ВЫВОДЫ ===============
add_heading(doc, "6  Выводы")
add_p(doc, "Все четыре численных метода сошлись к найденной аналитически точке "
      "x* = (1,8; −0,4), f* = 21,6. Положительная определённость матрицы Гессе "
      "(главные миноры 4 и 20) подтвердила, что эта точка — глобальный минимум.")
add_dash(doc, "Метод Флетчера–Ривса оказался самым быстрым ([[VAL:fr_iter]] итерации) благодаря свойству квадратичного завершения: для квадратичной функции n переменных метод сходится за n шагов.")
add_dash(doc, "Методы Хука–Дживса и наискорейшего спуска показали практически одинаковое число итераций ([[VAL:hj_iter]] и [[VAL:sd_iter]] соответственно), но Хук–Дживс не требует вычисления производных функции, что делает его более универсальным.")
add_dash(doc, "Метод наискорейшего спуска даёт характерную зигзагообразную сходимость: соседние направления антиградиента ортогональны, что замедляет сходимость на функциях с вытянутыми линиями уровня.")
add_dash(doc, "Метод Нелдера–Мида самый «дешёвый» по требованиям (не нужны ни производные, ни одномерный поиск), но самый медленный ([[VAL:nm_iter]] итерация). Он надёжен и подходит для негладких функций, но для квадратичных задач уступает градиентным методам.")
add_dash(doc, "Для квадратичных задач с известной матрицей Q наиболее эффективен метод сопряжённых градиентов Флетчера–Ривса.")

# ============= ПРИЛОЖЕНИЕ ===============
add_heading(doc, "Приложение А  Репозиторий с исходным кодом и ноутбуками")
add_p(doc, "Полный исходный код реализаций, ноутбуки Jupyter с прогоном каждого метода и "
      "построением графиков, а также скрипт сборки настоящего отчёта опубликованы в "
      "репозитории GitHub:")
add_p(doc, "https://github.com/megusto0/mp-1", align="center", indent_first=False)
add_p(doc, "Структура репозитория:")
add_dash(doc, "`src/target.py` — целевая функция, градиент, матрицы Q и b;")
add_dash(doc, "`src/nelder_mead.py` — реализация метода Нелдера–Мида;")
add_dash(doc, "`src/hooke_jeeves.py` — методы деления отрезка пополам и Хука–Дживса;")
add_dash(doc, "`src/gradient_methods.py` — методы наискорейшего спуска и Флетчера–Ривса;")
add_dash(doc, "`notebooks/01_analysis.ipynb` — аналитическое исследование (SymPy + графики);")
add_dash(doc, "`notebooks/02_nelder_mead.ipynb` — прогон и визуализация Нелдера–Мида;")
add_dash(doc, "`notebooks/03_hooke_jeeves.ipynb` — прогон и визуализация Хука–Дживса;")
add_dash(doc, "`notebooks/04_steepest_descent.ipynb` — прогон и визуализация SD;")
add_dash(doc, "`notebooks/05_fletcher_reeves.ipynb` — прогон и визуализация FR;")
add_dash(doc, "`notebooks/06_comparison.ipynb` — сравнительный анализ методов;")
add_dash(doc, "`build_report.py` — заполнение настоящего шаблона артефактами;")
add_dash(doc, "`requirements.txt` — `numpy sympy matplotlib pandas python-docx`.")
add_p(doc, "Каждый ноутбук открывается в Google Colab по соответствующему бейджу в README "
      "репозитория и выполняется без установки на локальную машину.")

# Сохранить
from pathlib import Path

out = Path(__file__).with_name("report_template.docx")
doc.save(out)
print(f"Saved: {out}")
