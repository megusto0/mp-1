"""build_report.py — наполнение шаблона lab1_v16 артефактами после прогона ноутбуков.

Запуск:  python build_report.py  (из корня репозитория)

Ожидаемая структура:
  report_template.docx
  values.json                  — словарь {ключ: значение} для меток [[VAL:key]]
  src/<file>.py                — исходники, из которых берутся функции для [[CODE:...]]
  figures/*.png                — рисунки для [[FIG:...]]
  tables/*.csv                 — таблицы для [[TBL:...]]
Результат: report_v16.docx

Метки:
  [[VAL:key]]                  — текстовая подстановка из values.json
  [[FIG:path|caption]]         — рисунок + подпись (по центру под рисунком)
  [[TBL:path|caption]]         — подпись + таблица (подпись слева над таблицей по ГОСТ)
  [[CODE:src/file.py|name]]    — листинг функции `name` из файла (моноширинно)
"""
from __future__ import annotations

import ast
import csv
import json
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

# -------- Конфиг путей --------
ROOT = Path(__file__).parent
TEMPLATE = ROOT / "report_template.docx"
OUTPUT = ROOT / "report_v16.docx"
VALUES_JSON = ROOT / "values.json"
FIGURES_DIR = ROOT / "figures"
TABLES_DIR = ROOT / "tables"

FONT = "Times New Roman"
MONO = "Consolas"

# Регулярки меток
RE_VAL = re.compile(r"\[\[VAL:([^\]]+)\]\]")
RE_FIG = re.compile(r"^\s*\[\[FIG:([^|]+)\|([^\]]+)\]\]\s*$")
RE_TBL = re.compile(r"^\s*\[\[TBL:([^|]+)\|([^\]]+)\]\]\s*$")
RE_CODE = re.compile(r"^\s*\[\[CODE:([^|]+)\|([^\]]+)\]\]\s*$")


def _set_run_font(run, *, mono=False, size=Pt(14), bold=False, italic=False):
    run.font.name = MONO if mono else FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    name = MONO if mono else FONT
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)


def _replace_paragraph_text(paragraph, new_text, *, mono=False, size=Pt(14)):
    # Полностью очищаем runs и вставляем новый
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    run = paragraph.add_run(new_text)
    _set_run_font(run, mono=mono, size=size)


def _extract_function_source(file_path: Path, func_name: str) -> str:
    """Достать тело функции `func_name` из Python-файла."""
    source = file_path.read_text(encoding="utf-8")
    tree = ast.parse(source)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == func_name:
            start = node.lineno - 1
            end = node.end_lineno
            lines = source.splitlines()[start:end]
            return "\n".join(lines)
    raise ValueError(f"Function {func_name} not found in {file_path}")


def _apply_inline_values(paragraph, values: dict) -> bool:
    """Заменить все [[VAL:key]] во всех runs абзаца. Возвращает True если что-то заменено."""
    full_text = "".join(r.text for r in paragraph.runs)
    if not RE_VAL.search(full_text):
        return False
    new_text = RE_VAL.sub(lambda m: str(values.get(m.group(1), f"<?{m.group(1)}?>")), full_text)
    _replace_paragraph_text(paragraph, new_text)
    return True


def _insert_image_paragraph(paragraph, image_path: Path, caption: str, fig_no: str):
    """В существующем абзаце заменить плейсхолдер на рисунок и подпись."""
    if not image_path.exists():
        _replace_paragraph_text(paragraph, f"[РИСУНОК НЕ НАЙДЕН: {image_path}]", mono=True)
        return
    # Очистить
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))
    # Подпись — отдельный абзац после
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph as PCls
    cap_para = PCls(new_p, paragraph._parent)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.paragraph_format.space_before = Pt(3)
    cap_para.paragraph_format.space_after = Pt(12)
    cap_run = cap_para.add_run(f"Рисунок {fig_no}. {caption}")
    _set_run_font(cap_run, size=Pt(14), italic=True)


def _insert_table(paragraph, csv_path: Path, caption: str, tab_no: str):
    """Заменить плейсхолдер: подпись сверху по ГОСТ + таблица из CSV."""
    if not csv_path.exists():
        _replace_paragraph_text(paragraph, f"[ТАБЛИЦА НЕ НАЙДЕНА: {csv_path}]", mono=True)
        return
    # Подпись (заменяет сам параграф)
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run(f"Таблица {tab_no}. {caption}")
    _set_run_font(run, size=Pt(14))

    # Прочитать CSV
    rows = []
    with csv_path.open(encoding="utf-8") as f:
        reader = csv.reader(f)
        for r in reader:
            rows.append(r)
    if not rows:
        return
    n_cols = len(rows[0])

    # Создать таблицу временно в конце документа, перенести после подписи
    doc = paragraph._parent
    while not hasattr(doc, "add_table"):
        doc = doc._parent
    # Body.add_table требует width; ширина текстовой области = 210 - 30 - 15 = 165 мм
    try:
        table = doc.add_table(rows=len(rows), cols=n_cols, width=Mm(165))
    except TypeError:
        table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(val)
            _set_run_font(run, size=Pt(13), bold=(i == 0))
    # Переместить таблицу сразу после подписи
    paragraph._p.addnext(table._tbl)


def _insert_code_block(paragraph, file_rel: str, func_name: str):
    src_path = ROOT / file_rel
    if not src_path.exists():
        _replace_paragraph_text(paragraph, f"[ИСХОДНИК НЕ НАЙДЕН: {file_rel}]", mono=True)
        return
    try:
        code = _extract_function_source(src_path, func_name)
    except ValueError as e:
        _replace_paragraph_text(paragraph, f"[{e}]", mono=True)
        return
    # Очистить и стилизовать как кодовый блок
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    # Серый фон через shading
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    # Каждая строка как run с break
    lines = code.splitlines()
    for i, line in enumerate(lines):
        run = paragraph.add_run(line)
        _set_run_font(run, mono=True, size=Pt(10))
        if i < len(lines) - 1:
            br = OxmlElement("w:br")
            run._r.append(br)


def fill_document():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    values = json.loads(VALUES_JSON.read_text(encoding="utf-8")) if VALUES_JSON.exists() else {}

    doc = Document(str(TEMPLATE))

    # Счётчики для нумерации
    fig_counter = {}
    tab_counter = {}

    # Обход всех абзацев документа
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs:
        text = "".join(r.text for r in paragraph.runs).strip()

        # CODE
        m = RE_CODE.match(text)
        if m:
            _insert_code_block(paragraph, m.group(1), m.group(2))
            continue

        # FIG
        m = RE_FIG.match(text)
        if m:
            img_path, caption = m.group(1), m.group(2)
            # Номер по контексту — извлекаем из имени файла fig_X_Y_*.png
            fname = Path(img_path).name
            mnum = re.match(r"fig_(\d+)_(\d+)_", fname)
            fig_no = f"{mnum.group(1)}.{mnum.group(2)}" if mnum else "X.Y"
            _insert_image_paragraph(paragraph, ROOT / img_path, caption, fig_no)
            continue

        # TBL
        m = RE_TBL.match(text)
        if m:
            csv_path, caption = m.group(1), m.group(2)
            fname = Path(csv_path).name
            mnum = re.match(r"tab_(\d+)_(\d+)_", fname)
            if not mnum:
                mnum = re.match(r"tab_(\d+)_", fname)
                tab_no = f"{mnum.group(1)}.1" if mnum else "X.Y"
            else:
                tab_no = f"{mnum.group(1)}.{mnum.group(2)}"
            _insert_table(paragraph, ROOT / csv_path, caption, tab_no)
            continue

        # VAL inline
        _apply_inline_values(paragraph, values)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    fill_document()
